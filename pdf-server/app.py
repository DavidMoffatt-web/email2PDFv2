#!/usr/bin/env python3
"""
Lightweight Python PDF Generation Server
Provides high-quality HTML-to-PDF conversion for Outlook add-in
"""

import os
import json
import base64
import tempfile
import logging
from datetime import datetime
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
import weasyprint
from weasyprint import HTML, CSS
from werkzeug.exceptions import BadRequest, InternalServerError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configuration
MAX_CONTENT_LENGTH = 50 * 1024 * 1024  # 50MB max request size
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'pdf-generator'
    })

@app.route('/convert', methods=['POST'])
def convert_to_pdf():
    """
    Convert HTML content to PDF
    
    Expected JSON payload:
    {
        "html": "HTML content",
        "css": "Optional CSS content",
        "options": {
            "page_size": "A4",
            "margin": "20mm",
            "orientation": "portrait"
        }
    }
    """
    try:
        # Validate request
        if not request.is_json:
            raise BadRequest("Content-Type must be application/json")
        
        data = request.get_json()
        if not data or 'html' not in data:
            raise BadRequest("Missing 'html' field in request body")
        
        html_content = data['html']
        css_content = data.get('css', '')
        options = data.get('options', {})
        
        logger.info(f"Processing PDF conversion request, HTML length: {len(html_content)}")
        
        # Default options
        page_size = options.get('page_size', 'A4')
        margin = options.get('margin', '20mm')
        orientation = options.get('orientation', 'portrait')
        
        # Enhanced CSS for better PDF rendering
        base_css = f"""
        @page {{
            size: {page_size} {orientation};
            margin: {margin};
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            margin: 0;
            padding: 0;
        }}
        
        /* Email content styling */
        .email-content {{
            max-width: 100%;
            word-wrap: break-word;
        }}
        
        /* Image handling */
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 10px 0;
        }}
        
        /* Table styling */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        
        table, th, td {{
            border: 1px solid #ddd;
        }}
        
        th, td {{
            padding: 8px;
            text-align: left;
        }}
        
        th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}
        
        /* Lists */
        ul, ol {{
            margin: 10px 0;
            padding-left: 20px;
        }}
        
        li {{
            margin: 5px 0;
        }}
        
        /* Headings */
        h1, h2, h3, h4, h5, h6 {{
            margin: 20px 0 10px 0;
            line-height: 1.2;
        }}
        
        /* Paragraphs */
        p {{
            margin: 10px 0;
        }}
        
        /* Links */
        a {{
            color: #0066cc;
            text-decoration: underline;
        }}
        
        /* Quote blocks */
        blockquote {{
            margin: 15px 0;
            padding: 10px 20px;
            border-left: 4px solid #ddd;
            background-color: #f9f9f9;
        }}
        
        /* Code blocks */
        pre, code {{
            font-family: 'Courier New', monospace;
            background-color: #f5f5f5;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        
        pre {{
            padding: 10px;
            overflow-x: auto;
        }}
        """
        
        # Combine CSS
        final_css = base_css + "\n" + css_content
        
        # Create temporary file for PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            try:
                # Generate PDF using WeasyPrint
                html_doc = HTML(string=html_content)
                css_doc = CSS(string=final_css)
                
                html_doc.write_pdf(
                    temp_pdf.name,
                    stylesheets=[css_doc],
                    optimize_images=True
                )
                
                # Read PDF content
                with open(temp_pdf.name, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Return PDF as base64
                pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
                
                logger.info(f"PDF generated successfully, size: {len(pdf_content)} bytes")
                
                return jsonify({
                    'success': True,
                    'pdf': pdf_base64,
                    'size': len(pdf_content),
                    'message': 'PDF generated successfully'
                })
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_pdf.name)
                except OSError:
                    pass
                    
    except BadRequest as e:
        logger.warning(f"Bad request: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Internal server error during PDF generation'
        }), 500

@app.route('/convert-binary', methods=['POST'])
def convert_to_pdf_binary():
    """
    Convert HTML content to PDF and return raw binary
    Same as /convert but returns PDF binary directly instead of base64 JSON
    """
    try:
        # Validate request
        if not request.is_json:
            raise BadRequest("Content-Type must be application/json")
        
        data = request.get_json()
        if not data or 'html' not in data:
            raise BadRequest("Missing 'html' field in request body")
        
        html_content = data['html']
        css_content = data.get('css', '')
        options = data.get('options', {})
        
        logger.info(f"Processing binary PDF conversion request, HTML length: {len(html_content)}")
        
        # Default options
        page_size = options.get('page_size', 'A4')
        margin = options.get('margin', '20mm')
        orientation = options.get('orientation', 'portrait')
        
        # Enhanced CSS with better styling
        base_css = f"""
        @page {{
            size: {page_size} {orientation};
            margin: {margin};
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 11pt;
        }}
        
        /* Images */
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 10px 0;
        }}
        
        /* Table styling */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        
        table, th, td {{
            border: 1px solid #ddd;
        }}
        
        th, td {{
            padding: 8px;
            text-align: left;
        }}
        
        th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}
        
        /* Lists */
        ul, ol {{
            margin: 10px 0;
            padding-left: 20px;
        }}
        
        li {{
            margin: 5px 0;
        }}
        
        /* Headings */
        h1, h2, h3, h4, h5, h6 {{
            margin: 20px 0 10px 0;
            line-height: 1.2;
        }}
        
        /* Paragraphs */
        p {{
            margin: 10px 0;
        }}
        
        /* Links */
        a {{
            color: #0066cc;
            text-decoration: underline;
        }}
        
        /* Code blocks */
        pre, code {{
            font-family: 'Courier New', monospace;
            background-color: #f5f5f5;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        
        pre {{
            padding: 10px;
            margin: 10px 0;
            border: 1px solid #ddd;
        }}
        
        /* Blockquotes */
        blockquote {{
            margin: 10px 0;
            padding: 10px 20px;
            border-left: 4px solid #ddd;
            background-color: #f9f9f9;
        }}
        """
        
        # Combine base CSS with custom CSS
        final_css = base_css + "\n" + css_content
        
        # Create temporary file for PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            try:
                # Generate PDF using WeasyPrint
                html_doc = HTML(string=html_content)
                css_doc = CSS(string=final_css)
                
                html_doc.write_pdf(
                    temp_pdf.name,
                    stylesheets=[css_doc],
                    optimize_images=True
                )
                
                # Read PDF content
                with open(temp_pdf.name, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                logger.info(f"Binary PDF generated successfully, size: {len(pdf_content)} bytes")
                
                # Return raw PDF binary
                return Response(
                    pdf_content,
                    mimetype='application/pdf',
                    headers={
                        'Content-Disposition': 'attachment; filename="email.pdf"',
                        'Content-Length': str(len(pdf_content))
                    }
                )
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_pdf.name)
                except OSError:
                    pass
                    
    except BadRequest as e:
        logger.warning(f"Bad request: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    
    except Exception as e:
        logger.error(f"Error generating binary PDF: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Internal server error during PDF generation'
        }), 500

@app.route('/convert-file', methods=['POST'])
def convert_to_pdf_file():
    """
    Convert HTML content to PDF and return as file download
    Same input as /convert but returns PDF file directly
    """
    try:
        # Validate request
        if not request.is_json:
            raise BadRequest("Content-Type must be application/json")
        
        data = request.get_json()
        if not data or 'html' not in data:
            raise BadRequest("Missing 'html' field in request body")
        
        html_content = data['html']
        css_content = data.get('css', '')
        options = data.get('options', {})
        filename = options.get('filename', 'email.pdf')
        
        logger.info(f"Processing PDF file conversion request")
        
        # Default options
        page_size = options.get('page_size', 'A4')
        margin = options.get('margin', '20mm')
        orientation = options.get('orientation', 'portrait')
        
        # Use same CSS as convert endpoint
        base_css = f"""
        @page {{
            size: {page_size} {orientation};
            margin: {margin};
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
        }}
        img {{ max-width: 100%; height: auto; }}
        table {{ width: 100%; border-collapse: collapse; }}
        table, th, td {{ border: 1px solid #ddd; }}
        th, td {{ padding: 8px; text-align: left; }}
        """
        
        final_css = base_css + "\n" + css_content
        
        # Create temporary file for PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            try:
                # Generate PDF
                html_doc = HTML(string=html_content)
                css_doc = CSS(string=final_css)
                
                html_doc.write_pdf(
                    temp_pdf.name,
                    stylesheets=[css_doc],
                    optimize_images=True
                )
                
                logger.info(f"PDF file generated successfully")
                
                return send_file(
                    temp_pdf.name,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/pdf'
                )
                
            except Exception as e:
                try:
                    os.unlink(temp_pdf.name)
                except OSError:
                    pass
                raise e
                
    except BadRequest as e:
        logger.warning(f"Bad request: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    
    except Exception as e:
        logger.error(f"Error generating PDF file: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Internal server error during PDF generation'
        }), 500

@app.route('/convert-with-attachments', methods=['POST'])
def convert_with_attachments():
    """
    Convert HTML content to PDF with attachment processing
    
    Expected JSON payload:
    {
        "html": "HTML content",
        "attachments": [
            {
                "name": "filename.docx",
                "content": "base64_encoded_content",
                "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        ],
        "mode": "full|individual", 
        "options": {
            "page_size": "A4",
            "margin": "20mm"
        }
    }
    """
    try:
        # Validate request
        if not request.is_json:
            raise BadRequest("Content-Type must be application/json")
        
        data = request.get_json()
        if not data or 'html' not in data:
            raise BadRequest("Missing 'html' field in request body")
        
        html_content = data['html']
        attachments = data.get('attachments', [])
        mode = data.get('mode', 'full')
        css_content = data.get('css', '')
        options = data.get('options', {})
        
        logger.info(f"Processing conversion with {len(attachments)} attachments, mode: {mode}")
        
        # Default options
        page_size = options.get('page_size', 'A4')
        margin = options.get('margin', '20mm')
        orientation = options.get('orientation', 'portrait')
        
        if mode == 'individual':
            # Return multiple PDFs as a ZIP file
            return handle_individual_pdfs(html_content, attachments, css_content, page_size, margin, orientation)
        else:
            # Merge everything into one PDF
            return handle_full_pdf(html_content, attachments, css_content, page_size, margin, orientation)
            
    except BadRequest as e:
        logger.warning(f"Bad request: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 400
    
    except Exception as e:
        logger.error(f"Error processing attachments: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Internal server error during attachment processing'
        }), 500

def handle_full_pdf(html_content, attachments, css_content, page_size, margin, orientation):
    """Handle full PDF mode - merge email and attachments into one PDF"""
    import io
    from PyPDF2 import PdfMerger
    
    try:
        merger = PdfMerger()
        
        # First, create PDF from email HTML
        email_pdf_content = create_pdf_from_html(html_content, css_content, page_size, margin, orientation)
        merger.append(io.BytesIO(email_pdf_content))
        
        # Process each attachment and convert to PDF
        for i, attachment in enumerate(attachments):
            try:
                logger.info(f"Processing attachment {i+1}: {attachment.get('name', 'unknown')}")
                attachment_pdf = convert_attachment_to_pdf(attachment, page_size, margin, orientation)
                if attachment_pdf:
                    merger.append(io.BytesIO(attachment_pdf))
            except Exception as e:
                logger.warning(f"Failed to convert attachment {attachment.get('name', 'unknown')}: {str(e)}")
                # Continue with other attachments
        
        # Write merged PDF to bytes
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        
        pdf_content = output.getvalue()
        output.close()
        
        logger.info(f"Merged PDF generated successfully, size: {len(pdf_content)} bytes")
        
        return Response(
            pdf_content,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': 'attachment; filename="email-with-attachments.pdf"',
                'Content-Length': str(len(pdf_content))
            }
        )
        
    except Exception as e:
        logger.error(f"Error creating merged PDF: {str(e)}")
        raise

def handle_individual_pdfs(html_content, attachments, css_content, page_size, margin, orientation):
    """Handle individual PDF mode - return ZIP with separate PDFs"""
    import zipfile
    import io
    
    try:
        # Create ZIP file in memory
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add email PDF
            email_pdf = create_pdf_from_html(html_content, css_content, page_size, margin, orientation)
            zip_file.writestr("email.pdf", email_pdf)
            
            # Add each attachment as separate PDF
            for i, attachment in enumerate(attachments):
                try:
                    attachment_name = attachment.get('name', f'attachment_{i+1}')
                    safe_name = "".join(c for c in attachment_name if c.isalnum() or c in "._- ").strip()
                    pdf_name = f"{safe_name}.pdf"
                    
                    logger.info(f"Converting attachment to PDF: {attachment_name}")
                    attachment_pdf = convert_attachment_to_pdf(attachment, page_size, margin, orientation)
                    
                    if attachment_pdf:
                        zip_file.writestr(pdf_name, attachment_pdf)
                    else:
                        # Create a placeholder PDF if conversion failed
                        placeholder_html = f"""
                        <html><body style="font-family: Arial, sans-serif; padding: 20px;">
                        <h2>Attachment: {attachment_name}</h2>
                        <p><strong>Type:</strong> {attachment.get('contentType', 'Unknown')}</p>
                        <p><strong>Status:</strong> Could not convert to PDF</p>
                        <p>This attachment type is not supported for PDF conversion.</p>
                        </body></html>
                        """
                        placeholder_pdf = create_pdf_from_html(placeholder_html, "", page_size, margin, orientation)
                        zip_file.writestr(pdf_name, placeholder_pdf)
                        
                except Exception as e:
                    logger.warning(f"Failed to process attachment {attachment.get('name', 'unknown')}: {str(e)}")
        
        zip_content = zip_buffer.getvalue()
        zip_buffer.close()
        
        logger.info(f"ZIP with individual PDFs created, size: {len(zip_content)} bytes")
        
        return Response(
            zip_content,
            mimetype='application/zip',
            headers={
                'Content-Disposition': 'attachment; filename="email-pdfs.zip"',
                'Content-Length': str(len(zip_content))
            }
        )
        
    except Exception as e:
        logger.error(f"Error creating individual PDFs: {str(e)}")
        raise

def create_pdf_from_html(html_content, css_content, page_size, margin, orientation):
    """Create PDF from HTML content"""
    # Enhanced CSS with better styling
    base_css = f"""
    @page {{
        size: {page_size} {orientation};
        margin: {margin};
    }}
    
    body {{
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        line-height: 1.6;
        color: #333;
        font-size: 11pt;
    }}
    
    img {{
        max-width: 100%;
        height: auto;
        display: block;
        margin: 10px 0;
    }}
    
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 15px 0;
    }}
    
    table, th, td {{
        border: 1px solid #ddd;
    }}
    
    th, td {{
        padding: 8px;
        text-align: left;
    }}
    
    th {{
        background-color: #f5f5f5;
        font-weight: bold;
    }}
    """
    
    # Combine base CSS with custom CSS
    final_css = base_css + "\n" + css_content
    
    # Create temporary file for PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
        try:
            # Generate PDF using WeasyPrint
            html_doc = HTML(string=html_content)
            css_doc = CSS(string=final_css)
            
            html_doc.write_pdf(
                temp_pdf.name,
                stylesheets=[css_doc],
                optimize_images=True
            )
            
            # Read PDF content
            with open(temp_pdf.name, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            return pdf_content
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_pdf.name)
            except OSError:
                pass

def convert_attachment_to_pdf(attachment, page_size, margin, orientation):
    """Convert various attachment types to PDF"""
    try:
        name = attachment.get('name', '')
        content_type = attachment.get('contentType', '')
        content_b64 = attachment.get('content', '')
        
        if not content_b64:
            logger.warning(f"No content provided for attachment: {name}")
            return None
        
        # Decode base64 content
        try:
            content_bytes = base64.b64decode(content_b64)
        except Exception as e:
            logger.error(f"Failed to decode base64 content for {name}: {str(e)}")
            return None
        
        logger.info(f"Converting {name} ({content_type}) to PDF")
        
        # Handle different file types
        if content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            return convert_docx_to_pdf(content_bytes, name, page_size, margin, orientation)
        elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
            return convert_excel_to_pdf(content_bytes, name, page_size, margin, orientation)
        elif content_type == 'application/pdf':
            return content_bytes  # Already PDF
        elif content_type.startswith('image/'):
            return convert_image_to_pdf(content_bytes, name, page_size, margin, orientation)
        elif content_type.startswith('text/'):
            return convert_text_to_pdf(content_bytes, name, page_size, margin, orientation)
        else:
            logger.warning(f"Unsupported content type for PDF conversion: {content_type}")
            return create_unsupported_file_pdf(name, content_type, page_size, margin, orientation)
            
    except Exception as e:
        logger.error(f"Error converting attachment {attachment.get('name', 'unknown')}: {str(e)}")
        return None

def convert_docx_to_pdf(content_bytes, name, page_size, margin, orientation):
    """Convert DOCX to PDF"""
    try:
        from docx import Document
        import io
        
        logger.info(f"Loading DOCX document: {name}, size: {len(content_bytes)} bytes")
        
        # Load DOCX document
        doc = Document(io.BytesIO(content_bytes))
        
        # Extract text and basic formatting
        html_content = f"""
        <html>
        <head>
            <title>{name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
                h1, h2, h3 {{ color: #333; }}
                p {{ margin: 10px 0; }}
                .document-header {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-left: 4px solid #007acc; 
                    margin-bottom: 20px; 
                }}
            </style>
        </head>
        <body>
        <div class="document-header">
            <h1>Document: {name}</h1>
        </div>
        """
        
        # Count paragraphs for debugging
        total_paragraphs = len(doc.paragraphs)
        content_paragraphs = 0
        
        logger.info(f"Document has {total_paragraphs} paragraphs")
        
        # Extract paragraphs with better handling
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                content_paragraphs += 1
                logger.debug(f"Paragraph {i+1}: '{para_text[:50]}...' (style: {paragraph.style.name})")
                
                # Handle different paragraph styles
                if paragraph.style.name.startswith('Heading'):
                    # Extract heading level
                    heading_parts = paragraph.style.name.split()
                    if len(heading_parts) > 1 and heading_parts[1].isdigit():
                        level = heading_parts[1]
                    else:
                        level = '1'
                    html_content += f"<h{level}>{para_text}</h{level}>\n"
                elif paragraph.style.name in ['Title', 'Subtitle']:
                    html_content += f"<h1>{para_text}</h1>\n"
                else:
                    # Regular paragraph
                    html_content += f"<p>{para_text}</p>\n"
        
        # If no content was found, try to extract in a different way
        if content_paragraphs == 0:
            logger.warning(f"No content paragraphs found in {name}, trying alternative extraction")
            
            # Try extracting from document body directly
            try:
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                
                if full_text.strip():
                    # Split into lines and create paragraphs
                    lines = [line.strip() for line in full_text.split('\n') if line.strip()]
                    for line in lines[:20]:  # Limit to first 20 lines
                        html_content += f"<p>{line}</p>\n"
                    logger.info(f"Extracted {len(lines)} lines of text using alternative method")
                else:
                    html_content += "<p><em>No text content could be extracted from this document.</em></p>\n"
                    logger.warning(f"No text content found in {name}")
            except Exception as alt_e:
                logger.error(f"Alternative extraction failed for {name}: {str(alt_e)}")
                html_content += "<p><em>Document could not be processed - it may be encrypted or corrupted.</em></p>\n"
        else:
            logger.info(f"Successfully extracted {content_paragraphs} paragraphs from {name}")
        
        html_content += "</body></html>"
        
        # Convert to PDF
        return create_pdf_from_html(html_content, "", page_size, margin, orientation)
        
    except Exception as e:
        logger.error(f"Error converting DOCX {name}: {str(e)}")
        return create_unsupported_file_pdf(name, "DOCX Document", page_size, margin, orientation)

def convert_excel_to_pdf(content_bytes, name, page_size, margin, orientation):
    """Convert Excel to PDF"""
    try:
        from openpyxl import load_workbook
        import io
        
        # Load Excel workbook
        wb = load_workbook(io.BytesIO(content_bytes), read_only=True)
        
        html_content = f"""
        <html>
        <head>
            <title>{name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
                h2 {{ color: #333; }}
            </style>
        </head>
        <body>
        <h1>Spreadsheet: {name}</h1>
        """
        
        # Process each worksheet
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            html_content += f"<h2>Sheet: {sheet_name}</h2>"
            html_content += "<table>"
            
            # Get data from sheet (limit to reasonable size)
            rows = list(ws.iter_rows(max_row=100, max_col=20, values_only=True))
            
            for i, row in enumerate(rows):
                if any(cell is not None for cell in row):
                    html_content += "<tr>"
                    for cell in row:
                        cell_value = str(cell) if cell is not None else ""
                        tag = "th" if i == 0 else "td"
                        html_content += f"<{tag}>{cell_value}</{tag}>"
                    html_content += "</tr>"
            
            html_content += "</table>"
        
        html_content += "</body></html>"
        
        # Convert to PDF
        return create_pdf_from_html(html_content, "", page_size, margin, orientation)
        
    except Exception as e:
        logger.error(f"Error converting Excel {name}: {str(e)}")
        return create_unsupported_file_pdf(name, "Excel Spreadsheet", page_size, margin, orientation)

def convert_image_to_pdf(content_bytes, name, page_size, margin, orientation):
    """Convert image to PDF"""
    try:
        import base64
        
        # Create data URL for the image
        image_b64 = base64.b64encode(content_bytes).decode()
        
        html_content = f"""
        <html>
        <head>
            <title>{name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; text-align: center; }}
                img {{ max-width: 100%; max-height: 80vh; }}
                h1 {{ color: #333; }}
            </style>
        </head>
        <body>
        <h1>Image: {name}</h1>
        <img src="data:image/png;base64,{image_b64}" alt="{name}">
        </body>
        </html>
        """
        
        return create_pdf_from_html(html_content, "", page_size, margin, orientation)
        
    except Exception as e:
        logger.error(f"Error converting image {name}: {str(e)}")
        return create_unsupported_file_pdf(name, "Image File", page_size, margin, orientation)

def convert_text_to_pdf(content_bytes, name, page_size, margin, orientation):
    """Convert text file to PDF"""
    try:
        # Try to decode as text
        try:
            text_content = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = content_bytes.decode('latin-1')
            except UnicodeDecodeError:
                text_content = str(content_bytes)
        
        # Convert line breaks to HTML
        text_content = text_content.replace('\n', '<br>')
        
        html_content = f"""
        <html>
        <head>
            <title>{name}</title>
            <style>
                body {{ font-family: monospace; margin: 20px; line-height: 1.4; }}
                h1 {{ color: #333; font-family: Arial, sans-serif; }}
                .content {{ white-space: pre-wrap; }}
            </style>
        </head>
        <body>
        <h1>Text File: {name}</h1>
        <div class="content">{text_content}</div>
        </body>
        </html>
        """
        
        return create_pdf_from_html(html_content, "", page_size, margin, orientation)
        
    except Exception as e:
        logger.error(f"Error converting text file {name}: {str(e)}")
        return create_unsupported_file_pdf(name, "Text File", page_size, margin, orientation)

def create_unsupported_file_pdf(name, file_type, page_size, margin, orientation):
    """Create a PDF for unsupported file types"""
    html_content = f"""
    <html>
    <head>
        <title>{name}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; text-align: center; }}
            .container {{ border: 2px dashed #ccc; padding: 40px; margin: 20px 0; }}
            h1 {{ color: #333; }}
            .file-icon {{ font-size: 48px; margin: 20px 0; }}
            .details {{ background-color: #f5f5f5; padding: 20px; margin: 20px 0; text-align: left; }}
        </style>
    </head>
    <body>
    <div class="container">
        <div class="file-icon">📄</div>
        <h1>Attachment: {name}</h1>
        <p><strong>File Type:</strong> {file_type}</p>
        <div class="details">
            <p><strong>Note:</strong> This file type cannot be automatically converted to PDF content.</p>
            <p>The original file was attached to the email and would need to be opened with appropriate software.</p>
        </div>
    </div>
    </body>
    </html>
    """
    
    return create_pdf_from_html(html_content, "", page_size, margin, orientation)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('DEBUG', 'false').lower() == 'true'
    
    logger.info(f"Starting PDF server on port {port}, debug={debug}")
    app.run(host='0.0.0.0', port=port, debug=debug)
