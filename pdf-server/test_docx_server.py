#!/usr/bin/env python3
"""
Simple Flask app for testing DOCX conversion without WeasyPrint
"""

import os
import logging
import io
import base64
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

def extract_docx_content(content_bytes, name):
    """Extract content from DOCX file and return as JSON"""
    try:
        logger.info(f"Loading DOCX document: {name}, size: {len(content_bytes)} bytes")
        
        # Load DOCX document
        doc = Document(io.BytesIO(content_bytes))
        
        # Count paragraphs for debugging
        total_paragraphs = len(doc.paragraphs)
        content_paragraphs = 0
        
        logger.info(f"Document has {total_paragraphs} paragraphs")
        
        extracted_content = []
        
        # Extract paragraphs with better handling
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                content_paragraphs += 1
                logger.debug(f"Paragraph {i+1}: '{para_text[:50]}...' (style: {paragraph.style.name})")
                
                extracted_content.append({
                    'index': i + 1,
                    'text': para_text,
                    'style': paragraph.style.name,
                    'type': 'heading' if paragraph.style.name.startswith('Heading') or paragraph.style.name in ['Title', 'Subtitle'] else 'paragraph'
                })
        
        # If no content was found, try alternative extraction
        if content_paragraphs == 0:
            logger.warning(f"No content paragraphs found in {name}, trying alternative extraction")
            
            try:
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                
                if full_text.strip():
                    lines = [line.strip() for line in full_text.split('\n') if line.strip()]
                    for i, line in enumerate(lines[:20]):  # Limit to first 20 lines
                        extracted_content.append({
                            'index': i + 1,
                            'text': line,
                            'style': 'Normal',
                            'type': 'paragraph'
                        })
                    logger.info(f"Extracted {len(lines)} lines of text using alternative method")
                else:
                    logger.warning(f"No text content found in {name}")
                    extracted_content.append({
                        'index': 1,
                        'text': 'No text content could be extracted from this document.',
                        'style': 'Normal',
                        'type': 'error'
                    })
            except Exception as alt_e:
                logger.error(f"Alternative extraction failed for {name}: {str(alt_e)}")
                extracted_content.append({
                    'index': 1,
                    'text': 'Document could not be processed - it may be encrypted or corrupted.',
                    'style': 'Normal',
                    'type': 'error'
                })
        else:
            logger.info(f"Successfully extracted {content_paragraphs} paragraphs from {name}")
        
        return {
            'success': True,
            'filename': name,
            'total_paragraphs': total_paragraphs,
            'content_paragraphs': content_paragraphs,
            'extracted_content': extracted_content
        }
        
    except Exception as e:
        logger.error(f"Error processing DOCX {name}: {str(e)}")
        return {
            'success': False,
            'filename': name,
            'error': str(e),
            'extracted_content': []
        }

@app.route('/test-docx', methods=['POST'])
def test_docx():
    """Test endpoint for DOCX content extraction"""
    try:
        data = request.json
        
        if not data or 'fileData' not in data or 'filename' not in data:
            return jsonify({'error': 'Missing fileData or filename'}), 400
        
        # Decode base64 file data
        file_data = base64.b64decode(data['fileData'])
        filename = data['filename']
        
        logger.info(f"Received DOCX file: {filename}, size: {len(file_data)} bytes")
        
        # Extract content
        result = extract_docx_content(file_data, filename)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in test-docx endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'message': 'DOCX test server is running'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    debug = os.environ.get('DEBUG', 'false').lower() == 'true'
    
    logger.info(f"Starting DOCX test server on port {port}, debug={debug}")
    app.run(host='0.0.0.0', port=port, debug=debug)
