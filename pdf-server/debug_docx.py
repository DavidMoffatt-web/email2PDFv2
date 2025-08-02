#!/usr/bin/env python3
"""
Test script to debug DOCX conversion on the server side
"""

import base64
import io
from docx import Document

def debug_docx_conversion():
    """Debug DOCX conversion with detailed logging"""
    
    # Create a test DOCX
    doc = Document()
    doc.add_heading('My Test Document', 0)
    doc.add_paragraph('This is the first paragraph with important content.')
    doc.add_paragraph('This is the second paragraph.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content under section 1.')
    doc.add_heading('Subsection 1.1', level=2)
    doc.add_paragraph('Content under subsection 1.1.')
    
    # Save to bytes (like it would come from client)
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    content_bytes = doc_bytes.read()
    
    print(f"Test DOCX created with {len(content_bytes)} bytes")
    
    # Convert to base64 (like it would come from client)
    base64_content = base64.b64encode(content_bytes).decode('utf-8')
    print(f"Base64 encoded: {len(base64_content)} characters")
    
    # Decode it back (like server would do)
    decoded_bytes = base64.b64decode(base64_content)
    print(f"Decoded back to {len(decoded_bytes)} bytes")
    
    # Now test the DOCX extraction (like server would do)
    doc_loaded = Document(io.BytesIO(decoded_bytes))
    
    print(f"\nDocument loaded successfully")
    print(f"Total paragraphs: {len(doc_loaded.paragraphs)}")
    
    html_content = """
    <html>
    <head>
        <title>Test Document</title>
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }
            h1, h2, h3 { color: #333; }
            p { margin: 10px 0; }
            .document-header { 
                background-color: #f8f9fa; 
                padding: 15px; 
                border-left: 4px solid #007acc; 
                margin-bottom: 20px; 
            }
        </style>
    </head>
    <body>
    <div class="document-header">
        <h1>Document: Test Document</h1>
    </div>
    """
    
    content_paragraphs = 0
    
    for i, paragraph in enumerate(doc_loaded.paragraphs):
        para_text = paragraph.text.strip()
        if para_text:
            content_paragraphs += 1
            print(f"Paragraph {i+1}: '{para_text}' (style: {paragraph.style.name})")
            
            # Handle different paragraph styles
            if paragraph.style.name.startswith('Heading'):
                # Extract heading level
                heading_parts = paragraph.style.name.split()
                if len(heading_parts) > 1 and heading_parts[1].isdigit():
                    level = heading_parts[1]
                else:
                    level = '1'
                html_content += f"<h{level}>{para_text}</h{level}>\n"
            elif paragraph.style.name in ['Title', 'Subtitle']:
                html_content += f"<h1>{para_text}</h1>\n"
            else:
                # Regular paragraph
                html_content += f"<p>{para_text}</p>\n"
    
    html_content += "</body></html>"
    
    print(f"\nSuccessfully extracted {content_paragraphs} paragraphs")
    print(f"\nGenerated HTML (length: {len(html_content)}):")
    print(html_content)
    
    return base64_content, html_content

if __name__ == "__main__":
    base64_content, html_content = debug_docx_conversion()
    
    # Save the base64 content to a file for testing
    with open('test_docx_base64.txt', 'w') as f:
        f.write(base64_content)
    
    print(f"\nBase64 content saved to test_docx_base64.txt")
    print(f"You can use this for testing the server endpoint")
