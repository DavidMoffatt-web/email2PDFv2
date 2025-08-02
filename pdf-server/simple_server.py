#!/usr/bin/env python3
"""
Simplified server for testing without WeasyPrint
"""

import os
import logging
import io
import base64
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

def convert_docx_to_html(content_bytes, name):
    """Convert DOCX to HTML (without PDF generation)"""
    try:
        logger.info(f"Loading DOCX document: {name}, size: {len(content_bytes)} bytes")
        
        # Load DOCX document
        doc = Document(io.BytesIO(content_bytes))
        
        # Extract text and basic formatting
        html_content = f"""
        <html>
        <head>
            <title>{name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }}
                h1, h2, h3 {{ color: #333; }}
                p {{ margin: 10px 0; }}
                .document-header {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-left: 4px solid #007acc; 
                    margin-bottom: 20px; 
                }}
            </style>
        </head>
        <body>
        <div class="document-header">
            <h1>Document: {name}</h1>
        </div>
        """
        
        # Count paragraphs for debugging
        total_paragraphs = len(doc.paragraphs)
        content_paragraphs = 0
        
        logger.info(f"Document has {total_paragraphs} paragraphs")
        
        # Extract paragraphs with better handling
        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            if para_text:
                content_paragraphs += 1
                logger.debug(f"Paragraph {i+1}: '{para_text[:50]}...' (style: {paragraph.style.name})")
                
                # Handle different paragraph styles
                if paragraph.style.name.startswith('Heading'):
                    # Extract heading level
                    heading_parts = paragraph.style.name.split()
                    if len(heading_parts) > 1 and heading_parts[1].isdigit():
                        level = heading_parts[1]
                    else:
                        level = '1'
                    html_content += f"<h{level}>{para_text}</h{level}>\n"
                elif paragraph.style.name in ['Title', 'Subtitle']:
                    html_content += f"<h1>{para_text}</h1>\n"
                else:
                    # Regular paragraph
                    html_content += f"<p>{para_text}</p>\n"
        
        # If no content was found, try to extract in a different way
        if content_paragraphs == 0:
            logger.warning(f"No content paragraphs found in {name}, trying alternative extraction")
            
            # Try extracting from document body directly
            try:
                full_text = ""
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
                
                if full_text.strip():
                    # Split into lines and create paragraphs
                    lines = [line.strip() for line in full_text.split('\n') if line.strip()]
                    for line in lines[:20]:  # Limit to first 20 lines
                        html_content += f"<p>{line}</p>\n"
                    logger.info(f"Extracted {len(lines)} lines of text using alternative method")
                else:
                    html_content += "<p><em>No text content could be extracted from this document.</em></p>\n"
                    logger.warning(f"No text content found in {name}")
            except Exception as alt_e:
                logger.error(f"Alternative extraction failed for {name}: {str(alt_e)}")
                html_content += "<p><em>Document could not be processed - it may be encrypted or corrupted.</em></p>\n"
        else:
            logger.info(f"Successfully extracted {content_paragraphs} paragraphs from {name}")
        
        html_content += "</body></html>"
        
        return {
            'success': True,
            'filename': name,
            'total_paragraphs': total_paragraphs,
            'content_paragraphs': content_paragraphs,
            'html_content': html_content
        }
        
    except Exception as e:
        logger.error(f"Error converting DOCX {name}: {str(e)}")
        return {
            'success': False,
            'filename': name,
            'error': str(e)
        }

@app.route('/convert-with-attachments', methods=['POST'])
def convert_with_attachments():
    """Simplified version that just processes attachments and returns HTML"""
    try:
        data = request.json
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        html = data.get('html', '')
        attachments = data.get('attachments', [])
        mode = data.get('mode', 'full')
        
        logger.info(f"Received request with {len(attachments)} attachments in {mode} mode")
        
        results = {
            'main_html': html,
            'attachments': []
        }
        
        # Process each attachment
        for attachment in attachments:
            name = attachment.get('name', 'unknown')
            content = attachment.get('content', '')
            content_type = attachment.get('contentType', 'application/octet-stream')
            
            logger.info(f"Processing attachment: {name} ({content_type})")
            
            if not content:
                logger.warning(f"No content for attachment: {name}")
                results['attachments'].append({
                    'name': name,
                    'status': 'no_content',
                    'error': 'No content provided'
                })
                continue
            
            try:
                # Decode base64 content
                content_bytes = base64.b64decode(content)
                logger.info(f"Decoded {name}: {len(content_bytes)} bytes")
                
                # Process DOCX files
                if name.lower().endswith('.docx'):
                    result = convert_docx_to_html(content_bytes, name)
                    results['attachments'].append({
                        'name': name,
                        'type': 'docx',
                        'status': 'processed',
                        'result': result
                    })
                else:
                    results['attachments'].append({
                        'name': name,
                        'type': 'other',
                        'status': 'not_processed',
                        'message': f'File type not supported for conversion: {content_type}'
                    })
                    
            except Exception as e:
                logger.error(f"Error processing attachment {name}: {str(e)}")
                results['attachments'].append({
                    'name': name,
                    'status': 'error',
                    'error': str(e)
                })
        
        return jsonify(results)
        
    except Exception as e:
        logger.error(f"Error in convert-with-attachments endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'message': 'Simplified server is running'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5002))
    debug = os.environ.get('DEBUG', 'false').lower() == 'true'
    
    logger.info(f"Starting simplified server on port {port}, debug={debug}")
    app.run(host='0.0.0.0', port=port, debug=debug)
