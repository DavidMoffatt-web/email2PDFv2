#!/usr/bin/env python3
"""
Test script to verify DOCX content extraction without WeasyPrint dependency
"""

import io
import base64
from docx import Document

def test_docx_extraction():
    """Test DOCX content extraction"""
    
    # Create a simple test DOCX in memory
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph with some content.')
    doc.add_paragraph('This is another paragraph to verify extraction.')
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('More content under section 2.')
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    content_bytes = doc_bytes.read()
    
    print(f"Created test DOCX with {len(content_bytes)} bytes")
    
    # Now test extraction
    doc = Document(io.BytesIO(content_bytes))
    
    total_paragraphs = len(doc.paragraphs)
    content_paragraphs = 0
    
    print(f"Document has {total_paragraphs} paragraphs")
    
    for i, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if para_text:
            content_paragraphs += 1
            print(f"Paragraph {i+1}: '{para_text[:50]}...' (style: {paragraph.style.name})")
    
    print(f"Found {content_paragraphs} paragraphs with content")
    
    # Test HTML generation
    html_content = """
    <html>
    <head>
        <title>Test Document</title>
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }
            h1, h2, h3 { color: #333; }
            p { margin: 10px 0; }
        </style>
    </head>
    <body>
    <h1>Document: Test Document</h1>
    """
    
    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if para_text:
            if paragraph.style.name.startswith('Heading'):
                heading_parts = paragraph.style.name.split()
                if len(heading_parts) > 1 and heading_parts[1].isdigit():
                    level = heading_parts[1]
                else:
                    level = '1'
                html_content += f"<h{level}>{para_text}</h{level}>\n"
            elif paragraph.style.name in ['Title', 'Subtitle']:
                html_content += f"<h1>{para_text}</h1>\n"
            else:
                html_content += f"<p>{para_text}</p>\n"
    
    html_content += "</body></html>"
    
    print("\nGenerated HTML:")
    print(html_content)
    
    return html_content

if __name__ == "__main__":
    test_docx_extraction()
